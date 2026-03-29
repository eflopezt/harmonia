"""
Vistas para gestión de contratos laborales, adendas, renovaciones y alertas.
"""
from datetime import date
from decimal import Decimal

from dateutil.relativedelta import relativedelta

from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.core.paginator import Paginator
from django.db.models import Q, Count, Case, When, Value, CharField
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.utils import timezone
from django.views.decorators.http import require_POST

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

from personal.models import Personal, Contrato, RenovacionContrato, Adenda, PlantillaContrato
from personal.forms import ContratoForm, AdendaForm, RenovacionContratoForm

solo_admin = user_passes_test(lambda u: u.is_superuser, login_url='login')

# ─── Estilos Excel ────────────────────────────────────────────────────────────
HEADER_FILL  = PatternFill("solid", fgColor="0D2B27")
HEADER_FONT  = Font(color="FFFFFF", bold=True, size=10)
ALT_FILL     = PatternFill("solid", fgColor="F0FDFA")
THIN_BORDER  = Border(
    left=Side(style='thin', color='CCCCCC'),
    right=Side(style='thin', color='CCCCCC'),
    top=Side(style='thin', color='CCCCCC'),
    bottom=Side(style='thin', color='CCCCCC'),
)

# ─── Umbrales de alerta ───────────────────────────────────────────────
DIAS_ALERTA_CONTRATO = [30, 15, 7]
DIAS_ALERTA_PRUEBA   = [15, 7]


# ═════════════════════════════════════════════════════════════════════════════
# PANEL PRINCIPAL
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contratos_panel(request):
    """Panel de contratos laborales con alertas y seguimiento."""
    from personal.models import Area
    hoy = timezone.localdate()

    activos = Personal.objects.filter(estado='Activo').select_related('subarea__area')

    # ── Filtros del panel ──────────────────────────────────────────────
    buscar = request.GET.get('buscar', '').strip()
    area_f = request.GET.get('area', '')
    tipo_f = request.GET.get('tipo', '')
    estado_f = request.GET.get('estado_contrato', '')
    fecha_desde = request.GET.get('fecha_desde', '').strip()
    fecha_hasta = request.GET.get('fecha_hasta', '').strip()

    if buscar:
        activos = activos.filter(
            Q(apellidos_nombres__icontains=buscar) | Q(nro_doc__icontains=buscar)
        )
    if area_f:
        activos = activos.filter(subarea__area_id=area_f)
    if tipo_f:
        activos = activos.filter(tipo_contrato=tipo_f)
    if estado_f == 'VIGENTE':
        activos = activos.filter(fecha_fin_contrato__gte=hoy)
    elif estado_f == 'VENCIDO':
        activos = activos.filter(fecha_fin_contrato__isnull=False, fecha_fin_contrato__lt=hoy)
    elif estado_f == 'INDEFINIDO':
        activos = activos.filter(tipo_contrato='INDEFINIDO')
    elif estado_f == 'SIN_DATOS':
        activos = activos.filter(tipo_contrato='')
    if fecha_desde:
        try:
            activos = activos.filter(fecha_fin_contrato__gte=date.fromisoformat(fecha_desde))
        except ValueError:
            pass
    if fecha_hasta:
        try:
            activos = activos.filter(fecha_fin_contrato__lte=date.fromisoformat(fecha_hasta))
        except ValueError:
            pass

    areas = Area.objects.filter(activa=True).order_by('nombre')

    # ── Contratos por vencer ──────────────────────────────────────────
    vencen_30 = activos.filter(
        fecha_fin_contrato__isnull=False,
        fecha_fin_contrato__gte=hoy,
        fecha_fin_contrato__lte=hoy + relativedelta(days=30),
    ).order_by('fecha_fin_contrato')

    vencen_60 = activos.filter(
        fecha_fin_contrato__isnull=False,
        fecha_fin_contrato__gt=hoy + relativedelta(days=30),
        fecha_fin_contrato__lte=hoy + relativedelta(days=60),
    ).order_by('fecha_fin_contrato')

    vencen_90 = activos.filter(
        fecha_fin_contrato__isnull=False,
        fecha_fin_contrato__gt=hoy + relativedelta(days=60),
        fecha_fin_contrato__lte=hoy + relativedelta(days=90),
    ).order_by('fecha_fin_contrato')

    vencidos = activos.filter(
        fecha_fin_contrato__isnull=False,
        fecha_fin_contrato__lt=hoy,
    ).order_by('fecha_fin_contrato')

    sin_contrato = activos.filter(tipo_contrato='').count()

    # ── Período de prueba activo ──────────────────────────────────────
    desde_max = hoy - relativedelta(months=12)
    candidatos_prueba = activos.filter(fecha_alta__gte=desde_max)

    en_prueba = []
    for p in candidatos_prueba:
        fin_prueba = p.fecha_fin_periodo_prueba
        if fin_prueba and fin_prueba >= hoy:
            dias_restantes = (fin_prueba - hoy).days
            en_prueba.append({
                'personal': p,
                'fin_prueba': fin_prueba,
                'dias_restantes': dias_restantes,
                'alerta': dias_restantes <= 15,
            })
    en_prueba.sort(key=lambda x: x['fin_prueba'])

    # ── Estadísticas por tipo de contrato ────────────────────────────
    por_tipo = (
        activos
        .exclude(tipo_contrato='')
        .values('tipo_contrato')
        .annotate(total=Count('id'))
        .order_by('-total')
    )
    sin_tipo = activos.filter(tipo_contrato='').count()

    # ── KPIs ──────────────────────────────────────────────────────────
    total_activos = activos.count()
    indefinidos   = activos.filter(tipo_contrato='INDEFINIDO').count()
    plazo_fijo    = activos.filter(tipo_contrato__in=[
        'PLAZO_FIJO', 'INICIO_ACTIVIDAD', 'NECESIDAD_MERCADO',
        'RECONVERSION_EMPRESARIAL', 'OBRA_SERVICIO', 'DISCONTINUO',
        'TEMPORADA', 'SUPLENCIA', 'EMERGENCIA',
    ]).count()

    # ── Últimas adendas y renovaciones (modelo nuevo) ────────────────
    ultimas_adendas = Adenda.objects.select_related(
        'contrato__personal',
    ).order_by('-creado_en')[:5]

    ultimas_renovaciones = RenovacionContrato.objects.select_related(
        'contrato_original__personal',
    ).order_by('-creado_en')[:5]

    context = {
        'hoy': hoy,
        'total_activos': total_activos,
        'indefinidos': indefinidos,
        'plazo_fijo': plazo_fijo,
        'sin_tipo': sin_tipo,
        'sin_contrato': sin_contrato,
        'vencen_30': vencen_30,
        'vencen_60': vencen_60,
        'vencen_90': vencen_90,
        'vencidos': vencidos,
        'en_prueba': en_prueba,
        'por_tipo': por_tipo,
        'ultimas_adendas': ultimas_adendas,
        'ultimas_renovaciones': ultimas_renovaciones,
        'tab_activo': request.GET.get('tab', 'vencimientos'),
        # Filtros
        'buscar': buscar,
        'area_f': area_f,
        'tipo_f': tipo_f,
        'estado_f': estado_f,
        'fecha_desde': fecha_desde,
        'fecha_hasta': fecha_hasta,
        'areas': areas,
        'tipos': Personal.TIPO_CONTRATO_CHOICES,
    }
    return render(request, 'personal/contratos_panel.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# LISTA COMPLETA DE CONTRATOS
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contratos_lista(request):
    """Lista de todos los empleados con sus datos de contrato."""
    hoy = timezone.localdate()

    qs = Personal.objects.filter(estado='Activo').select_related('subarea__area').order_by('apellidos_nombres')

    # Filtros
    buscar = request.GET.get('buscar', '').strip()
    tipo_f = request.GET.get('tipo', '')
    estado_f = request.GET.get('estado_contrato', '')
    area_f = request.GET.get('area', '')
    vencimiento_f = request.GET.get('vencimiento', '')  # 30, 60, 90

    if buscar:
        qs = qs.filter(
            Q(apellidos_nombres__icontains=buscar) | Q(nro_doc__icontains=buscar)
        )
    if tipo_f:
        qs = qs.filter(tipo_contrato=tipo_f)
    if area_f:
        qs = qs.filter(subarea__area_id=area_f)

    if estado_f == 'VIGENTE':
        qs = qs.filter(fecha_fin_contrato__gte=hoy)
    elif estado_f == 'VENCIDO':
        qs = qs.filter(fecha_fin_contrato__lt=hoy)
    elif estado_f == 'INDEFINIDO':
        qs = qs.filter(tipo_contrato='INDEFINIDO')
    elif estado_f == 'SIN_DATOS':
        qs = qs.filter(tipo_contrato='')

    if vencimiento_f:
        try:
            dias_lim = int(vencimiento_f)
            qs = qs.filter(
                fecha_fin_contrato__isnull=False,
                fecha_fin_contrato__gte=hoy,
                fecha_fin_contrato__lte=hoy + relativedelta(days=dias_lim),
            )
        except ValueError:
            pass

    paginator = Paginator(qs, 30)
    page_obj = paginator.get_page(request.GET.get('page'))

    # Prefetch contrato vigente IDs for checkbox selection
    personal_ids = [p.pk for p in page_obj]
    contratos_vigentes = {}
    for c in Contrato.objects.filter(
        personal_id__in=personal_ids, estado='VIGENTE'
    ).values('personal_id', 'id').order_by('personal_id', '-fecha_inicio'):
        if c['personal_id'] not in contratos_vigentes:
            contratos_vigentes[c['personal_id']] = c['id']

    empleados = []
    for p in page_obj:
        dias = None
        estado_contrato = 'INDEFINIDO' if p.tipo_contrato == 'INDEFINIDO' else ''
        if p.fecha_fin_contrato:
            dias = (p.fecha_fin_contrato - hoy).days
            if dias < 0:
                estado_contrato = 'VENCIDO'
            elif dias <= 30:
                estado_contrato = 'URGENTE'
            else:
                estado_contrato = 'VIGENTE'
        elif not p.tipo_contrato:
            estado_contrato = 'SIN_DATOS'
        empleados.append({
            'personal': p,
            'dias_restantes': dias,
            'estado_contrato': estado_contrato,
            'contrato_id': contratos_vigentes.get(p.pk),
        })

    from personal.models import Area
    areas = Area.objects.filter(activa=True).order_by('nombre')

    context = {
        'empleados': empleados,
        'page_obj': page_obj,
        'buscar': buscar,
        'tipo_f': tipo_f,
        'estado_f': estado_f,
        'area_f': area_f,
        'vencimiento_f': vencimiento_f,
        'tipos': Personal.TIPO_CONTRATO_CHOICES,
        'areas': areas,
        'hoy': hoy,
        'total_resultados': paginator.count,
    }
    return render(request, 'personal/contratos_lista.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# DETALLE DE CONTRATO DE UN EMPLEADO
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_detalle(request, pk):
    """Detalle del contrato de un empleado con timeline de adendas y renovaciones."""
    personal = get_object_or_404(Personal, pk=pk)
    hoy = timezone.localdate()

    # Contratos del modelo nuevo (si existen)
    contratos = Contrato.objects.filter(
        personal=personal
    ).prefetch_related('adendas', 'renovaciones_salientes__contrato_nuevo').order_by('-fecha_inicio')

    # Timeline: combinar adendas y renovaciones ordenadas por fecha
    timeline = []
    for c in contratos:
        timeline.append({
            'tipo': 'contrato',
            'fecha': c.fecha_inicio,
            'titulo': f'Contrato {c.get_tipo_contrato_display()}',
            'detalle': f'{c.fecha_inicio.strftime("%d/%m/%Y")} — {c.fecha_fin.strftime("%d/%m/%Y") if c.fecha_fin else "Indefinido"}',
            'estado': c.estado,
            'objeto': c,
        })
        for adenda in c.adendas.all():
            timeline.append({
                'tipo': 'adenda',
                'fecha': adenda.fecha,
                'titulo': f'Adenda: {adenda.get_tipo_modificacion_display()}',
                'detalle': adenda.detalle[:100],
                'estado': '',
                'objeto': adenda,
            })
        for ren in c.renovaciones_salientes.all():
            timeline.append({
                'tipo': 'renovacion',
                'fecha': ren.fecha_renovacion,
                'titulo': 'Renovación de contrato',
                'detalle': ren.motivo[:100] if ren.motivo else 'Sin motivo registrado',
                'estado': '',
                'objeto': ren,
            })
    timeline.sort(key=lambda x: x['fecha'], reverse=True)

    # Info del contrato actual (campos en Personal)
    dias_restantes = None
    estado_contrato = 'INDEFINIDO' if personal.tipo_contrato == 'INDEFINIDO' else ''
    if personal.fecha_fin_contrato:
        dias_restantes = (personal.fecha_fin_contrato - hoy).days
        if dias_restantes < 0:
            estado_contrato = 'VENCIDO'
        elif dias_restantes <= 30:
            estado_contrato = 'URGENTE'
        else:
            estado_contrato = 'VIGENTE'
    elif not personal.tipo_contrato:
        estado_contrato = 'SIN_DATOS'

    contrato_vigente = contratos.filter(estado='VIGENTE').first()

    context = {
        'personal': personal,
        'contratos': contratos,
        'contrato_vigente': contrato_vigente,
        'timeline': timeline,
        'dias_restantes': dias_restantes,
        'estado_contrato': estado_contrato,
        'hoy': hoy,
    }
    return render(request, 'personal/contrato_detalle.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# EDITAR CONTRATO (legacy — campos en Personal)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_editar(request, pk):
    """Edita los campos de contrato de un empleado."""
    personal = get_object_or_404(Personal, pk=pk)

    if request.method == 'POST':
        personal.tipo_contrato         = request.POST.get('tipo_contrato', '')
        personal.renovacion_automatica = bool(request.POST.get('renovacion_automatica'))
        personal.observaciones_contrato = request.POST.get('observaciones_contrato', '')

        fecha_inicio = request.POST.get('fecha_inicio_contrato', '').strip()
        fecha_fin    = request.POST.get('fecha_fin_contrato', '').strip()

        personal.fecha_inicio_contrato = date.fromisoformat(fecha_inicio) if fecha_inicio else None
        personal.fecha_fin_contrato    = date.fromisoformat(fecha_fin) if fecha_fin else None

        personal.save(update_fields=[
            'tipo_contrato', 'fecha_inicio_contrato', 'fecha_fin_contrato',
            'renovacion_automatica', 'observaciones_contrato',
        ])
        messages.success(request, f"Contrato de {personal.apellidos_nombres} actualizado.")
        return redirect('contrato_detalle', pk=personal.pk)

    context = {
        'personal': personal,
        'tipos': Personal.TIPO_CONTRATO_CHOICES,
    }
    return render(request, 'personal/contrato_editar.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# CREAR CONTRATO (modelo nuevo)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_crear(request, personal_pk):
    """Crea un nuevo contrato para un empleado usando el modelo Contrato."""
    personal = get_object_or_404(Personal, pk=personal_pk)

    if request.method == 'POST':
        form = ContratoForm(request.POST, request.FILES)
        if form.is_valid():
            contrato = form.save(commit=False)
            contrato.personal = personal
            contrato.registrado_por = request.user
            contrato.save()
            # Sincronizar con Personal
            contrato.sincronizar_con_personal()
            messages.success(request, f"Contrato creado para {personal.apellidos_nombres}.")
            return redirect('contrato_detalle', pk=personal.pk)
    else:
        initial = {
            'tipo_contrato': personal.tipo_contrato,
            'fecha_inicio': personal.fecha_inicio_contrato,
            'fecha_fin': personal.fecha_fin_contrato,
            'sueldo_pactado': personal.sueldo_base,
            'cargo_contrato': personal.cargo,
            'renovacion_automatica': personal.renovacion_automatica,
        }
        form = ContratoForm(initial=initial)

    plantillas = PlantillaContrato.objects.filter(activo=True).order_by('nombre')
    context = {
        'personal': personal,
        'form': form,
        'es_nuevo': True,
        'plantillas': plantillas,
    }
    return render(request, 'personal/contrato_form.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# EDITAR CONTRATO (modelo nuevo)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_editar_obj(request, pk):
    """Edita un contrato del modelo Contrato."""
    contrato = get_object_or_404(Contrato, pk=pk)
    personal = contrato.personal

    if request.method == 'POST':
        form = ContratoForm(request.POST, request.FILES, instance=contrato)
        if form.is_valid():
            form.save()
            if contrato.estado == 'VIGENTE':
                contrato.sincronizar_con_personal()
            messages.success(request, "Contrato actualizado.")
            return redirect('contrato_detalle', pk=personal.pk)
    else:
        form = ContratoForm(instance=contrato)

    context = {
        'personal': personal,
        'form': form,
        'contrato': contrato,
        'es_nuevo': False,
    }
    return render(request, 'personal/contrato_form.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# RENOVAR CONTRATO
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_renovar(request, pk):
    """Renueva un contrato existente creando uno nuevo y vinculándolos."""
    contrato_original = get_object_or_404(Contrato, pk=pk)
    personal = contrato_original.personal

    if request.method == 'POST':
        form = RenovacionContratoForm(request.POST)
        if form.is_valid():
            # Marcar original como renovado
            contrato_original.estado = 'RENOVADO'
            contrato_original.save(update_fields=['estado'])

            # Crear nuevo contrato
            nuevo = Contrato.objects.create(
                personal=personal,
                tipo_contrato=form.cleaned_data['tipo_contrato'],
                fecha_inicio=form.cleaned_data['fecha_inicio'],
                fecha_fin=form.cleaned_data['fecha_fin'],
                sueldo_pactado=form.cleaned_data.get('sueldo_pactado') or contrato_original.sueldo_pactado,
                cargo_contrato=contrato_original.cargo_contrato,
                jornada_semanal=contrato_original.jornada_semanal,
                renovacion_automatica=contrato_original.renovacion_automatica,
                observaciones=form.cleaned_data.get('observaciones', ''),
                estado='VIGENTE',
                registrado_por=request.user,
            )

            # Registrar renovación
            RenovacionContrato.objects.create(
                contrato_original=contrato_original,
                contrato_nuevo=nuevo,
                fecha_renovacion=timezone.localdate(),
                motivo=form.cleaned_data.get('motivo', ''),
                registrado_por=request.user,
            )

            # Sincronizar con Personal
            nuevo.sincronizar_con_personal()

            messages.success(request, f"Contrato de {personal.apellidos_nombres} renovado exitosamente.")
            return redirect('contrato_detalle', pk=personal.pk)
    else:
        initial = {
            'tipo_contrato': contrato_original.tipo_contrato,
            'fecha_inicio': contrato_original.fecha_fin + relativedelta(days=1) if contrato_original.fecha_fin else timezone.localdate(),
            'sueldo_pactado': contrato_original.sueldo_pactado,
        }
        form = RenovacionContratoForm(initial=initial)

    context = {
        'personal': personal,
        'contrato_original': contrato_original,
        'form': form,
    }
    return render(request, 'personal/contrato_renovar.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# ADENDAS
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def adenda_crear(request, contrato_pk):
    """Crea una adenda para un contrato."""
    contrato = get_object_or_404(Contrato, pk=contrato_pk)
    personal = contrato.personal

    if request.method == 'POST':
        form = AdendaForm(request.POST, request.FILES)
        if form.is_valid():
            adenda = form.save(commit=False)
            adenda.contrato = contrato
            adenda.registrado_por = request.user
            adenda.save()
            messages.success(request, "Adenda registrada exitosamente.")
            return redirect('contrato_detalle', pk=personal.pk)
    else:
        form = AdendaForm(initial={'fecha': timezone.localdate()})

    context = {
        'personal': personal,
        'contrato': contrato,
        'form': form,
    }
    return render(request, 'personal/adenda_form.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# EXPORTAR CONTRATOS A EXCEL
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contratos_exportar_excel(request):
    """Exporta la lista de contratos a Excel con filtros aplicados."""
    hoy = timezone.localdate()

    qs = Personal.objects.filter(estado='Activo').select_related('subarea__area').order_by('apellidos_nombres')

    # Aplicar mismos filtros que la lista
    buscar = request.GET.get('buscar', '').strip()
    tipo_f = request.GET.get('tipo', '')
    estado_f = request.GET.get('estado_contrato', '')
    area_f = request.GET.get('area', '')

    if buscar:
        qs = qs.filter(Q(apellidos_nombres__icontains=buscar) | Q(nro_doc__icontains=buscar))
    if tipo_f:
        qs = qs.filter(tipo_contrato=tipo_f)
    if area_f:
        qs = qs.filter(subarea__area_id=area_f)
    if estado_f == 'VIGENTE':
        qs = qs.filter(fecha_fin_contrato__gte=hoy)
    elif estado_f == 'VENCIDO':
        qs = qs.filter(fecha_fin_contrato__lt=hoy)
    elif estado_f == 'INDEFINIDO':
        qs = qs.filter(tipo_contrato='INDEFINIDO')
    elif estado_f == 'SIN_DATOS':
        qs = qs.filter(tipo_contrato='')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Contratos'

    # Título
    ws['A1'] = f'Reporte de Contratos Laborales — {hoy.strftime("%d/%m/%Y")}'
    ws['A1'].font = Font(size=13, bold=True, color='0D2B27')
    ws.merge_cells('A1:L1')
    ws['A1'].alignment = Alignment(horizontal='center')
    ws.row_dimensions[1].height = 28

    # Headers
    headers = [
        'N°', 'Trabajador', 'DNI', 'Cargo', 'Área', 'SubÁrea',
        'Modalidad', 'Inicio', 'Vencimiento', 'Estado', 'Días Restantes',
        'Sueldo Base',
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = THIN_BORDER

    RED_FILL   = PatternFill("solid", fgColor="FEE2E2")
    AMBER_FILL = PatternFill("solid", fgColor="FEF3C7")
    GREEN_FILL = PatternFill("solid", fgColor="D1FAE5")

    for i, p in enumerate(qs, 1):
        row = i + 2
        dias = None
        estado = ''
        if p.fecha_fin_contrato:
            dias = (p.fecha_fin_contrato - hoy).days
            if dias < 0:
                estado = 'VENCIDO'
            elif dias <= 30:
                estado = 'URGENTE'
            else:
                estado = 'VIGENTE'
        elif p.tipo_contrato == 'INDEFINIDO':
            estado = 'INDEFINIDO'
        elif not p.tipo_contrato:
            estado = 'SIN DATOS'

        # Color urgencia
        if estado == 'VENCIDO':
            fill = RED_FILL
        elif estado == 'URGENTE':
            fill = AMBER_FILL
        elif estado == 'VIGENTE':
            fill = GREEN_FILL
        else:
            fill = ALT_FILL if i % 2 == 0 else PatternFill()

        area_n = p.subarea.area.nombre if p.subarea and p.subarea.area else '—'
        subarea_n = p.subarea.nombre if p.subarea else '—'

        values = [
            i,
            p.apellidos_nombres,
            p.nro_doc,
            p.cargo or '—',
            area_n,
            subarea_n,
            p.get_tipo_contrato_display() or '—',
            p.fecha_inicio_contrato,
            p.fecha_fin_contrato,
            estado,
            dias if dias is not None else '—',
            float(p.sueldo_base or 0),
        ]

        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill = fill
            cell.border = THIN_BORDER
            cell.alignment = Alignment(
                horizontal='center' if col in (1, 3, 8, 9, 10, 11) else 'left',
                vertical='center'
            )
            if col == 3:
                cell.number_format = '@'
            elif col == 12:
                cell.number_format = '"S/ "#,##0.00'

    # Auto-width
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=10)
        from openpyxl.utils import get_column_letter
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 3, 50)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="contratos_{hoy.strftime("%Y%m%d")}.xlsx"'
    wb.save(response)
    return response


# ═════════════════════════════════════════════════════════════════════════════
# API: datos para dashboard de contratos (AJAX)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contratos_api_stats(request):
    """JSON con KPIs de contratos para gráficos."""
    hoy = timezone.localdate()
    activos = Personal.objects.filter(estado='Activo')

    por_tipo = {}
    for val, lbl in Personal.TIPO_CONTRATO_CHOICES:
        c = activos.filter(tipo_contrato=val).count()
        if c:
            por_tipo[lbl] = c

    return JsonResponse({
        'por_tipo': por_tipo,
        'vencen_7': activos.filter(
            fecha_fin_contrato__gte=hoy,
            fecha_fin_contrato__lte=hoy + relativedelta(days=7)
        ).count(),
        'vencen_30': activos.filter(
            fecha_fin_contrato__gte=hoy,
            fecha_fin_contrato__lte=hoy + relativedelta(days=30)
        ).count(),
        'vencen_60': activos.filter(
            fecha_fin_contrato__gte=hoy,
            fecha_fin_contrato__lte=hoy + relativedelta(days=60)
        ).count(),
        'vencen_90': activos.filter(
            fecha_fin_contrato__gte=hoy,
            fecha_fin_contrato__lte=hoy + relativedelta(days=90)
        ).count(),
        'vencidos': activos.filter(fecha_fin_contrato__lt=hoy).count(),
        'sin_contrato': activos.filter(tipo_contrato='').count(),
    })


# ═════════════════════════════════════════════════════════════════════════════
# ALERTAS DE CONTRATOS (widget para dashboard)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contratos_alertas_json(request):
    """JSON con alertas de contratos por vencer para widgets del dashboard."""
    hoy = timezone.localdate()
    activos = Personal.objects.filter(estado='Activo')

    alertas = []

    # Contratos vencidos
    vencidos = activos.filter(
        fecha_fin_contrato__isnull=False,
        fecha_fin_contrato__lt=hoy,
    ).select_related('subarea__area')
    for p in vencidos[:10]:
        dias = (hoy - p.fecha_fin_contrato).days
        alertas.append({
            'nivel': 'danger',
            'mensaje': f'{p.apellidos_nombres} — Contrato vencido hace {dias} día(s)',
            'personal_id': p.pk,
        })

    # Contratos por vencer en 30 días
    por_vencer = activos.filter(
        fecha_fin_contrato__gte=hoy,
        fecha_fin_contrato__lte=hoy + relativedelta(days=30),
    ).select_related('subarea__area').order_by('fecha_fin_contrato')
    for p in por_vencer[:10]:
        dias = (p.fecha_fin_contrato - hoy).days
        nivel = 'danger' if dias <= 7 else ('warning' if dias <= 15 else 'info')
        alertas.append({
            'nivel': nivel,
            'mensaje': f'{p.apellidos_nombres} — Contrato vence en {dias} día(s)',
            'personal_id': p.pk,
        })

    # Personal sin contrato
    sin_contrato_count = activos.filter(tipo_contrato='').count()
    if sin_contrato_count:
        alertas.append({
            'nivel': 'secondary',
            'mensaje': f'{sin_contrato_count} empleado(s) sin datos de contrato',
            'personal_id': None,
        })

    return JsonResponse({'alertas': alertas})


# ═════════════════════════════════════════════════════════════════════════════
# GENERAR PDF DE CONTRATO
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_generar_pdf(request, pk):
    """Genera un PDF profesional del contrato laboral."""
    from io import BytesIO
    from django.template.loader import get_template
    from xhtml2pdf import pisa
    from empresas.models import Empresa

    contrato = get_object_or_404(Contrato, pk=pk)
    personal = contrato.personal

    # Obtener empresa (del personal o la principal)
    empresa = personal.empresa
    if not empresa:
        empresa = Empresa.objects.filter(es_principal=True).first()
    if not empresa:
        empresa = Empresa.objects.first()

    # Convertir sueldo a letras (simplificado)
    sueldo_letras = ''
    if contrato.sueldo_pactado:
        try:
            from num2words import num2words
            sueldo_letras = num2words(float(contrato.sueldo_pactado), lang='es') + ' soles'
        except Exception:
            sueldo_letras = f'{contrato.sueldo_pactado} soles'

    context = {
        'contrato': contrato,
        'personal': personal,
        'empresa': empresa,
        'tipo_contrato_display': contrato.get_tipo_contrato_display(),
        'sueldo_letras': sueldo_letras,
    }

    template = get_template('personal/contrato_pdf.html')
    html = template.render(context)

    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode('utf-8')), result, encoding='utf-8')

    if pdf.err:
        return HttpResponse('Error generando el PDF', status=500)

    response = HttpResponse(result.getvalue(), content_type='application/pdf')
    nombre_archivo = f'contrato_{personal.nro_doc}_{contrato.fecha_inicio.strftime("%Y%m%d")}.pdf'
    # Si se pide descarga directa
    if request.GET.get('download') == '1':
        response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    else:
        response['Content-Disposition'] = f'inline; filename="{nombre_archivo}"'
    return response


# ═════════════════════════════════════════════════════════════════════════════
# IMPORTAR PLANTILLA DE CONTRATO (DOCX/PDF)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_importar_plantilla(request, pk):
    """Permite subir un archivo DOCX o PDF como plantilla adjunta al contrato."""
    contrato = get_object_or_404(Contrato, pk=pk)
    personal = contrato.personal

    if request.method == 'POST':
        archivo = request.FILES.get('archivo')
        if not archivo:
            messages.error(request, 'Debe seleccionar un archivo.')
            return redirect('contrato_detalle', pk=personal.pk)

        ext = archivo.name.lower().rsplit('.', 1)[-1] if '.' in archivo.name else ''
        if ext not in ('pdf', 'docx', 'doc'):
            messages.error(request, 'Solo se permiten archivos PDF o DOCX.')
            return redirect('contrato_detalle', pk=personal.pk)

        contrato.archivo_pdf = archivo
        contrato.save(update_fields=['archivo_pdf'])
        messages.success(request, f'Archivo "{archivo.name}" adjuntado al contrato.')
        return redirect('contrato_detalle', pk=personal.pk)

    context = {
        'personal': personal,
        'contrato': contrato,
    }
    return render(request, 'personal/contrato_importar_plantilla.html', context)


# ═════════════════════════════════════════════════════════════════════════════
# ANÁLISIS IA DE CONTRATO
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def contrato_analizar_ia(request, pk):
    """Analiza un contrato con IA: resumen, cláusulas clave, riesgos."""
    contrato = get_object_or_404(Contrato, pk=pk)
    personal = contrato.personal

    if not contrato.archivo_pdf:
        return JsonResponse({
            'ok': False,
            'error': 'El contrato no tiene un archivo PDF/DOCX adjunto para analizar.',
        })

    # Extraer texto del archivo
    texto = ''
    archivo_path = contrato.archivo_pdf.path
    ext = archivo_path.lower().rsplit('.', 1)[-1]

    try:
        if ext == 'pdf':
            try:
                from pdfminer.high_level import extract_text
                texto = extract_text(archivo_path)
            except ImportError:
                import subprocess
                # Fallback: intentar con xhtml2pdf o lectura básica
                texto = '[No se pudo extraer texto del PDF — instale pdfminer.six]'
        elif ext in ('docx', 'doc'):
            try:
                from docx import Document
                doc = Document(archivo_path)
                texto = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            except ImportError:
                texto = '[No se pudo extraer texto del DOCX — instale python-docx]'
    except Exception as e:
        return JsonResponse({'ok': False, 'error': f'Error extrayendo texto: {str(e)}'})

    if not texto or texto.startswith('[No se pudo'):
        return JsonResponse({
            'ok': False,
            'error': texto or 'No se pudo extraer texto del documento.',
        })

    # Enviar a IA para análisis
    from asistencia.services.ai_service import get_service
    svc = get_service()
    if not svc:
        return JsonResponse({
            'ok': False,
            'error': 'No hay servicio de IA configurado. Configure uno en Sistema > Configuracion > IA.',
        })

    prompt = (
        'Eres un abogado laboralista peruano experto en derecho laboral.\n'
        'Analiza el siguiente contrato de trabajo y proporciona:\n\n'
        '1. **RESUMEN**: Un resumen conciso del contrato (3-5 lineas).\n'
        '2. **CLAUSULAS CLAVE**: Lista las clausulas mas importantes.\n'
        '3. **OBSERVACIONES Y RIESGOS**: Identifica posibles problemas, '
        'clausulas abusivas o riesgos para el empleador o trabajador.\n'
        '4. **CUMPLIMIENTO LEGAL**: Verifica el cumplimiento con la legislacion '
        'laboral peruana (D.Leg. 728, D.S. 003-97-TR) e indica si falta algo.\n\n'
        'Responde en espanol. Se conciso pero completo.\n\n'
        '--- TEXTO DEL CONTRATO ---\n'
        f'{texto[:8000]}'  # Limitar para no exceder tokens
    )

    system = (
        'Eres un asistente juridico especializado en derecho laboral peruano. '
        'Analizas contratos de trabajo y proporcionas observaciones legales precisas.'
    )

    try:
        # Use generate with higher token limit
        resultado = svc.generate(prompt, system=system)
        if not resultado:
            return JsonResponse({
                'ok': False,
                'error': 'La IA no devolvio resultado. Intente nuevamente.',
            })

        return JsonResponse({
            'ok': True,
            'analisis': resultado,
            'contrato_id': contrato.pk,
            'trabajador': personal.apellidos_nombres,
        })
    except Exception as e:
        return JsonResponse({
            'ok': False,
            'error': f'Error al comunicarse con la IA: {str(e)}',
        })


# ═════════════════════════════════════════════════════════════════════════════
# ENVIAR CONTRATO POR EMAIL (Individual)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
@require_POST
def contrato_enviar_email(request, pk):
    """Genera el PDF del contrato y lo envía por email al trabajador."""
    from io import BytesIO
    from django.template.loader import get_template
    from django.core.mail import EmailMessage, get_connection
    from xhtml2pdf import pisa
    from empresas.models import Empresa

    contrato = get_object_or_404(Contrato, pk=pk)
    personal = contrato.personal

    # Verificar que el trabajador tiene email
    email_destino = (personal.correo_corporativo or
                     getattr(personal, 'correo_personal', '') or '')
    if not email_destino:
        messages.error(request, f'{personal.apellidos_nombres} no tiene correo registrado.')
        return redirect('contrato_detalle', pk=personal.pk)

    # Obtener empresa
    empresa = personal.empresa
    if not empresa:
        empresa = Empresa.objects.filter(es_principal=True).first()
    if not empresa:
        empresa = Empresa.objects.first()

    # Verificar SMTP
    if not empresa or not empresa.tiene_email_configurado:
        messages.error(request, 'No hay configuracion SMTP. Configure el correo en la empresa.')
        return redirect('contrato_detalle', pk=personal.pk)

    # Generar PDF
    sueldo_letras = ''
    if contrato.sueldo_pactado:
        try:
            from num2words import num2words
            sueldo_letras = num2words(float(contrato.sueldo_pactado), lang='es') + ' soles'
        except Exception:
            sueldo_letras = f'{contrato.sueldo_pactado} soles'

    context = {
        'contrato': contrato,
        'personal': personal,
        'empresa': empresa,
        'tipo_contrato_display': contrato.get_tipo_contrato_display(),
        'sueldo_letras': sueldo_letras,
    }

    template = get_template('personal/contrato_pdf.html')
    html = template.render(context)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode('utf-8')), result, encoding='utf-8')

    if pdf.err:
        messages.error(request, 'Error generando el PDF del contrato.')
        return redirect('contrato_detalle', pk=personal.pk)

    pdf_content = result.getvalue()
    nombre_pdf = f'contrato_{personal.nro_doc}_{contrato.fecha_inicio.strftime("%Y%m%d")}.pdf'

    # Enviar email
    smtp = empresa.get_smtp_config()
    try:
        connection = get_connection(
            host=smtp['host'],
            port=smtp['port'],
            username=smtp['username'],
            password=smtp['password'],
            use_tls=smtp['use_tls'],
            use_ssl=smtp['use_ssl'],
            fail_silently=False,
        )

        tipo_display = contrato.get_tipo_contrato_display()
        email = EmailMessage(
            subject=f'Contrato de Trabajo — {empresa.nombre_display}',
            body=(
                f'<p>Estimado(a) <strong>{personal.apellidos_nombres}</strong>,</p>'
                f'<p>Adjunto encontrara su contrato de trabajo bajo la modalidad '
                f'<strong>{tipo_display}</strong>.</p>'
                f'<p>Datos del contrato:</p>'
                f'<ul>'
                f'<li>Inicio: {contrato.fecha_inicio.strftime("%d/%m/%Y")}</li>'
                f'<li>Fin: {contrato.fecha_fin.strftime("%d/%m/%Y") if contrato.fecha_fin else "Indefinido"}</li>'
                f'<li>Cargo: {contrato.cargo_contrato or personal.cargo}</li>'
                f'</ul>'
                f'<p>Por favor, revise el documento y comuniquese con RRHH ante cualquier consulta.</p>'
                f'<p>Atentamente,<br><strong>{empresa.nombre_display}</strong></p>'
            ),
            from_email=smtp['from_email'],
            to=[email_destino],
            reply_to=[smtp['reply_to']] if smtp.get('reply_to') else None,
            connection=connection,
        )
        email.content_subtype = 'html'
        email.attach(nombre_pdf, pdf_content, 'application/pdf')
        email.send()

        messages.success(
            request,
            f'Contrato enviado exitosamente a {email_destino}.'
        )
    except Exception as e:
        messages.error(request, f'Error enviando email: {str(e)[:200]}')

    return redirect('contrato_detalle', pk=personal.pk)


# ═════════════════════════════════════════════════════════════════════════════
# ENVÍO MASIVO DE CONTRATOS POR EMAIL
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
@require_POST
def contratos_envio_masivo(request):
    """Envía múltiples contratos por email (uno a cada trabajador)."""
    from io import BytesIO
    from django.template.loader import get_template
    from django.core.mail import EmailMessage, get_connection
    from xhtml2pdf import pisa
    from empresas.models import Empresa

    contrato_ids = request.POST.getlist('contrato_ids')
    if not contrato_ids:
        messages.warning(request, 'No se seleccionaron contratos para enviar.')
        return redirect('contratos_lista')

    contratos = Contrato.objects.filter(pk__in=contrato_ids).select_related(
        'personal__empresa', 'personal__subarea__area'
    )

    if not contratos.exists():
        messages.error(request, 'No se encontraron los contratos seleccionados.')
        return redirect('contratos_lista')

    # Obtener empresa default
    empresa_default = Empresa.objects.filter(es_principal=True).first()
    if not empresa_default:
        empresa_default = Empresa.objects.first()

    enviados = 0
    errores = []

    for contrato in contratos:
        personal = contrato.personal
        empresa = personal.empresa or empresa_default

        if not empresa or not empresa.tiene_email_configurado:
            errores.append(f'{personal.apellidos_nombres}: Sin config SMTP')
            continue

        email_destino = (personal.correo_corporativo or
                         getattr(personal, 'correo_personal', '') or '')
        if not email_destino:
            errores.append(f'{personal.apellidos_nombres}: Sin email')
            continue

        # Generar PDF
        sueldo_letras = ''
        if contrato.sueldo_pactado:
            try:
                from num2words import num2words
                sueldo_letras = num2words(float(contrato.sueldo_pactado), lang='es') + ' soles'
            except Exception:
                sueldo_letras = f'{contrato.sueldo_pactado} soles'

        context = {
            'contrato': contrato,
            'personal': personal,
            'empresa': empresa,
            'tipo_contrato_display': contrato.get_tipo_contrato_display(),
            'sueldo_letras': sueldo_letras,
        }

        template = get_template('personal/contrato_pdf.html')
        html = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode('utf-8')), result, encoding='utf-8')

        if pdf.err:
            errores.append(f'{personal.apellidos_nombres}: Error generando PDF')
            continue

        pdf_content = result.getvalue()
        nombre_pdf = f'contrato_{personal.nro_doc}_{contrato.fecha_inicio.strftime("%Y%m%d")}.pdf'

        # Enviar
        smtp = empresa.get_smtp_config()
        try:
            connection = get_connection(
                host=smtp['host'],
                port=smtp['port'],
                username=smtp['username'],
                password=smtp['password'],
                use_tls=smtp['use_tls'],
                use_ssl=smtp['use_ssl'],
                fail_silently=False,
            )

            tipo_display = contrato.get_tipo_contrato_display()
            email = EmailMessage(
                subject=f'Contrato de Trabajo — {empresa.nombre_display}',
                body=(
                    f'<p>Estimado(a) <strong>{personal.apellidos_nombres}</strong>,</p>'
                    f'<p>Adjunto encontrara su contrato de trabajo '
                    f'(<strong>{tipo_display}</strong>).</p>'
                    f'<p>Atentamente,<br><strong>{empresa.nombre_display}</strong></p>'
                ),
                from_email=smtp['from_email'],
                to=[email_destino],
                reply_to=[smtp['reply_to']] if smtp.get('reply_to') else None,
                connection=connection,
            )
            email.content_subtype = 'html'
            email.attach(nombre_pdf, pdf_content, 'application/pdf')
            email.send()
            enviados += 1
        except Exception as e:
            errores.append(f'{personal.apellidos_nombres}: {str(e)[:100]}')

    if enviados:
        messages.success(request, f'{enviados} contrato(s) enviado(s) exitosamente por email.')
    if errores:
        msg_errores = '; '.join(errores[:5])
        if len(errores) > 5:
            msg_errores += f' ... y {len(errores) - 5} mas'
        messages.warning(request, f'Errores en envio: {msg_errores}')

    return redirect('contratos_lista')


# ═════════════════════════════════════════════════════════════════════════════
# PLANTILLAS DE CONTRATO (CRUD)
# ═════════════════════════════════════════════════════════════════════════════

@solo_admin
def plantilla_contrato_lista(request):
    """Lista todas las plantillas de contrato."""
    plantillas = PlantillaContrato.objects.all().order_by('-activo', 'nombre')
    context = {
        'plantillas': plantillas,
    }
    return render(request, 'personal/plantilla_contrato_lista.html', context)


@solo_admin
def plantilla_contrato_crear(request):
    """Crea una nueva plantilla de contrato."""
    if request.method == 'POST':
        nombre = request.POST.get('nombre', '').strip()
        tipo_contrato = request.POST.get('tipo_contrato', '')
        contenido_html = request.POST.get('contenido_html', '')

        if not nombre:
            messages.error(request, 'El nombre de la plantilla es obligatorio.')
            return render(request, 'personal/plantilla_contrato_form.html', {
                'tipos': Personal.TIPO_CONTRATO_CHOICES,
                'es_nuevo': True,
            })

        PlantillaContrato.objects.create(
            nombre=nombre,
            tipo_contrato=tipo_contrato,
            contenido_html=contenido_html,
        )
        messages.success(request, f'Plantilla "{nombre}" creada exitosamente.')
        return redirect('plantilla_contrato_lista')

    context = {
        'tipos': Personal.TIPO_CONTRATO_CHOICES,
        'es_nuevo': True,
    }
    return render(request, 'personal/plantilla_contrato_form.html', context)


@solo_admin
def plantilla_contrato_editar(request, pk):
    """Edita una plantilla de contrato existente."""
    plantilla = get_object_or_404(PlantillaContrato, pk=pk)

    if request.method == 'POST':
        plantilla.nombre = request.POST.get('nombre', '').strip()
        plantilla.tipo_contrato = request.POST.get('tipo_contrato', '')
        plantilla.contenido_html = request.POST.get('contenido_html', '')
        plantilla.activo = bool(request.POST.get('activo'))

        if not plantilla.nombre:
            messages.error(request, 'El nombre de la plantilla es obligatorio.')
        else:
            plantilla.save()
            messages.success(request, f'Plantilla "{plantilla.nombre}" actualizada.')
            return redirect('plantilla_contrato_lista')

    context = {
        'plantilla': plantilla,
        'tipos': Personal.TIPO_CONTRATO_CHOICES,
        'es_nuevo': False,
    }
    return render(request, 'personal/plantilla_contrato_form.html', context)


@solo_admin
@require_POST
def plantilla_contrato_eliminar(request, pk):
    """Elimina una plantilla de contrato."""
    plantilla = get_object_or_404(PlantillaContrato, pk=pk)
    nombre = plantilla.nombre
    plantilla.delete()
    messages.success(request, f'Plantilla "{nombre}" eliminada.')
    return redirect('plantilla_contrato_lista')


@solo_admin
def plantilla_contrato_importar(request):
    """Importa un archivo DOCX como plantilla de contrato."""
    if request.method == 'POST':
        archivo = request.FILES.get('archivo')
        nombre = request.POST.get('nombre', '').strip()
        tipo_contrato = request.POST.get('tipo_contrato', '')

        if not archivo:
            messages.error(request, 'Debe seleccionar un archivo.')
            return redirect('plantilla_contrato_importar')

        ext = archivo.name.lower().rsplit('.', 1)[-1] if '.' in archivo.name else ''
        if ext not in ('docx',):
            messages.error(request, 'Solo se permiten archivos DOCX para importar como plantilla.')
            return redirect('plantilla_contrato_importar')

        if not nombre:
            nombre = archivo.name.rsplit('.', 1)[0]

        # Extraer texto del DOCX
        contenido_html = ''
        try:
            from docx import Document
            doc = Document(archivo)
            parrafos = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parrafos.append(f'<p>{p.text}</p>')
            contenido_html = '\n'.join(parrafos)
        except ImportError:
            messages.error(request, 'python-docx no esta instalado. Instale con: pip install python-docx')
            return redirect('plantilla_contrato_importar')
        except Exception as e:
            messages.error(request, f'Error al procesar el archivo: {str(e)[:200]}')
            return redirect('plantilla_contrato_importar')

        PlantillaContrato.objects.create(
            nombre=nombre,
            tipo_contrato=tipo_contrato,
            contenido_html=contenido_html,
        )
        messages.success(request, f'Plantilla "{nombre}" importada exitosamente desde DOCX.')
        return redirect('plantilla_contrato_lista')

    context = {
        'tipos': Personal.TIPO_CONTRATO_CHOICES,
    }
    return render(request, 'personal/plantilla_contrato_importar.html', context)


# =========================================================================
# DETALLE DE PLANTILLA CON PREVIEW Y ANALISIS IA
# =========================================================================

@solo_admin
def plantilla_contrato_detalle(request, pk):
    """Muestra preview de la plantilla con datos de ejemplo y panel de IA."""
    import re
    plantilla = get_object_or_404(PlantillaContrato, pk=pk)

    # Datos de ejemplo para preview
    ejemplo = {
        'nombre_empleado': 'GARCIA LOPEZ, JUAN CARLOS',
        'cargo': 'Analista de Sistemas',
        'fecha_inicio': '01/04/2026',
        'fecha_fin': '30/09/2026',
        'remuneracion': 'S/ 4,500.00',
        'dni': '72345678',
        'empresa': 'ANDES MINING S.A.C.',
        'ruc_empresa': '20123456789',
        'direccion_empresa': 'Av. Industrial 234, Lima',
    }

    # Reemplazar placeholders
    preview_html = plantilla.contenido_html
    for key, val in ejemplo.items():
        preview_html = preview_html.replace('{{' + key + '}}', val)

    # Detectar placeholders usados
    placeholders = sorted(set(re.findall(r'\{\{(\w+)\}\}', plantilla.contenido_html)))

    context = {
        'plantilla': plantilla,
        'preview_html': preview_html,
        'placeholders': placeholders,
    }
    return render(request, 'personal/plantilla_contrato_detalle.html', context)


@solo_admin
def plantilla_contrato_ia(request, pk):
    """API: Chat con IA sobre una plantilla de contrato."""
    import json as _json
    plantilla = get_object_or_404(PlantillaContrato, pk=pk)

    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'POST requerido'}, status=405)

    try:
        data = _json.loads(request.body)
    except (ValueError, TypeError):
        return JsonResponse({'ok': False, 'error': 'JSON invalido'}, status=400)

    pregunta = data.get('pregunta', '').strip()
    historial = data.get('historial', [])
    if not pregunta:
        return JsonResponse({'ok': False, 'error': 'Pregunta vacia'})

    # Obtener servicio IA
    from asistencia.services.ai_service import get_service
    svc = get_service()
    if not svc:
        return JsonResponse({
            'ok': False,
            'error': 'No hay servicio de IA configurado. Configure uno en Sistema > Configuracion > IA.',
        })

    # Construir prompt con contexto de la plantilla
    contenido_plantilla = plantilla.contenido_html[:6000]
    system = (
        'Eres un asistente experto en contratos laborales peruanos y diseno de plantillas HTML. '
        'El usuario esta trabajando con una plantilla de contrato llamada "' + plantilla.nombre + '". '
        'La plantilla usa placeholders como {{nombre_empleado}}, {{cargo}}, {{fecha_inicio}}, etc. '
        'Responde en espanol. Se conciso pero util.\n\n'
        'Si el usuario pide cambios al HTML, incluye el HTML completo actualizado dentro de '
        'etiquetas [HTML_INICIO] y [HTML_FIN] para que el sistema pueda extraerlo.'
    )

    # Construir mensajes con historial
    prompt_parts = []
    prompt_parts.append('=== PLANTILLA HTML ACTUAL ===\n' + contenido_plantilla + '\n=== FIN PLANTILLA ===\n')

    for msg in historial[-4:]:
        role = msg.get('role', 'user')
        content = msg.get('content', '')
        if role == 'user':
            prompt_parts.append('USUARIO: ' + content)
        else:
            prompt_parts.append('ASISTENTE: ' + content[:500])

    prompt_parts.append('USUARIO: ' + pregunta)
    prompt = '\n\n'.join(prompt_parts)

    try:
        resultado = svc.generate(prompt, system=system)
        if not resultado:
            return JsonResponse({'ok': False, 'error': 'La IA no devolvio resultado.'})

        # Extraer HTML sugerido si existe
        html_sugerido = None
        if '[HTML_INICIO]' in resultado and '[HTML_FIN]' in resultado:
            import re
            match = re.search(r'\[HTML_INICIO\](.*?)\[HTML_FIN\]', resultado, re.DOTALL)
            if match:
                html_sugerido = match.group(1).strip()
                # Limpiar del texto de respuesta
                resultado = resultado[:resultado.index('[HTML_INICIO]')].strip()
                if not resultado:
                    resultado = 'He generado una version actualizada de la plantilla. Puedes ver el preview o aplicar los cambios.'

        response_data = {
            'ok': True,
            'respuesta': resultado,
        }
        if html_sugerido:
            response_data['html_sugerido'] = html_sugerido

        return JsonResponse(response_data)
    except Exception as e:
        return JsonResponse({'ok': False, 'error': f'Error de IA: {str(e)[:200]}'})


@solo_admin
@require_POST
def plantilla_contrato_aplicar_ia(request, pk):
    """Aplica el HTML sugerido por la IA a la plantilla."""
    import json as _json
    import re
    plantilla = get_object_or_404(PlantillaContrato, pk=pk)

    try:
        data = _json.loads(request.body)
    except (ValueError, TypeError):
        return JsonResponse({'ok': False, 'error': 'JSON invalido'}, status=400)

    html_nuevo = data.get('html', '').strip()
    if not html_nuevo:
        return JsonResponse({'ok': False, 'error': 'HTML vacio'})

    # Guardar
    plantilla.contenido_html = html_nuevo
    plantilla.save(update_fields=['contenido_html'])

    # Generar preview con datos de ejemplo
    ejemplo = {
        'nombre_empleado': 'GARCIA LOPEZ, JUAN CARLOS',
        'cargo': 'Analista de Sistemas',
        'fecha_inicio': '01/04/2026',
        'fecha_fin': '30/09/2026',
        'remuneracion': 'S/ 4,500.00',
        'dni': '72345678',
        'empresa': 'ANDES MINING S.A.C.',
        'ruc_empresa': '20123456789',
        'direccion_empresa': 'Av. Industrial 234, Lima',
    }
    preview_html = html_nuevo
    for key, val in ejemplo.items():
        preview_html = preview_html.replace('{{' + key + '}}', val)

    return JsonResponse({
        'ok': True,
        'preview_html': preview_html,
    })
