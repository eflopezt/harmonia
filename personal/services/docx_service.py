"""
Servicio de generación de documentos DOCX para contratos laborales.

Genera Word editables a partir del HTML resuelto de PlantillaContrato,
usando python-docx.
"""
import re
from io import BytesIO
from html.parser import HTMLParser

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


class SimpleHTMLToDocx(HTMLParser):
    """Parser simple que convierte HTML básico de contratos a python-docx paragraphs."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.is_bold = False
        self.is_italic = False
        self.is_underline = False
        self.in_list = False
        self.list_counter = 0
        self.in_title = False
        self.tag_stack = []
        self.skip_content = False

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        self.tag_stack.append(tag)

        if tag in ('h1', 'h2', 'h3', 'h4'):
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.is_bold = True
            self.in_title = True
        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            style = attrs_dict.get('style', '')
            if 'text-align:center' in style or 'text-align: center' in style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align:right' in style or 'text-align: right' in style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif tag == 'strong' or tag == 'b':
            self.is_bold = True
        elif tag == 'em' or tag == 'i':
            self.is_italic = True
        elif tag == 'u':
            self.is_underline = True
        elif tag == 'ol':
            self.in_list = True
            self.list_counter = 0
        elif tag == 'ul':
            self.in_list = True
            self.list_counter = -1  # Bullets
        elif tag == 'li':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if self.list_counter >= 0:
                self.list_counter += 1
                prefix = f"{self.list_counter}. "
            else:
                prefix = "• "
            run = self.current_paragraph.add_run(prefix)
            run.font.size = Pt(11)
            # Indent
            self.current_paragraph.paragraph_format.left_indent = Cm(1)
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
        elif tag == 'img':
            # Try to add image
            src = attrs_dict.get('src', '')
            if src.startswith('file:///'):
                path = src[8:]  # Remove file:///
                try:
                    if self.current_paragraph is None:
                        self.current_paragraph = self.doc.add_paragraph()
                    self.current_paragraph.add_run().add_picture(path, height=Inches(0.8))
                except Exception:
                    pass
        elif tag == 'div':
            style = attrs_dict.get('style', '')
            cls = attrs_dict.get('class', '')
            if 'clausula-titulo' in cls or 'clausula_titulo' in cls:
                self.current_paragraph = self.doc.add_paragraph()
                self.is_bold = True
            elif self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
        elif tag == 'table':
            pass  # Tables handled separately
        elif tag == 'hr':
            p = self.doc.add_paragraph()
            p.add_run('_' * 50)

    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

        if tag in ('h1', 'h2', 'h3', 'h4'):
            self.is_bold = False
            self.in_title = False
            self.current_paragraph = None
        elif tag == 'p':
            self.current_paragraph = None
        elif tag == 'strong' or tag == 'b':
            self.is_bold = False
        elif tag == 'em' or tag == 'i':
            self.is_italic = False
        elif tag == 'u':
            self.is_underline = False
        elif tag in ('ol', 'ul'):
            self.in_list = False
        elif tag == 'div':
            if self.is_bold:
                self.is_bold = False
            self.current_paragraph = None

    def handle_data(self, data):
        if self.skip_content:
            return
        text = data.strip()
        if not text:
            # Preserve single spaces
            if data and ' ' in data and self.current_paragraph:
                run = self.current_paragraph.add_run(' ')
                run.font.size = Pt(11)
            return

        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = self.current_paragraph.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.bold = self.is_bold
        run.italic = self.is_italic
        run.underline = self.is_underline

        if self.in_title:
            if 'h1' in self.tag_stack:
                run.font.size = Pt(14)
            elif 'h2' in self.tag_stack:
                run.font.size = Pt(13)
            elif 'h3' in self.tag_stack:
                run.font.size = Pt(12)


def generar_contrato_docx(html_resuelto, empresa=None, personal=None, contrato=None):
    """
    Genera un documento DOCX a partir del HTML resuelto de un contrato.

    Args:
        html_resuelto: HTML con placeholders ya resueltos
        empresa: instancia de Empresa
        personal: instancia de Personal
        contrato: instancia de Contrato

    Returns:
        BytesIO con el contenido del DOCX
    """
    doc = Document()

    # Configure page size (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ── Header with logo ──
    if empresa:
        header_section = section.header
        header_para = header_section.paragraphs[0] if header_section.paragraphs else header_section.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if empresa.logo:
            try:
                run = header_para.add_run()
                run.add_picture(empresa.logo.path, height=Inches(0.7))
            except Exception:
                pass

    # ── Parse HTML content ──
    # Clean HTML: remove <style>...</style>, <head>...</head>, etc.
    clean_html = re.sub(r'<style[^>]*>.*?</style>', '', html_resuelto, flags=re.DOTALL)
    clean_html = re.sub(r'<head[^>]*>.*?</head>', '', clean_html, flags=re.DOTALL)
    clean_html = re.sub(r'</?html[^>]*>', '', clean_html)
    clean_html = re.sub(r'</?body[^>]*>', '', clean_html)

    parser = SimpleHTMLToDocx(doc)
    parser.feed(clean_html)

    # ── Signature block ──
    doc.add_paragraph()  # Space before signatures

    # Two-column signature using a table
    table = doc.add_table(rows=3, cols=3)
    table.autofit = True

    # Employer signature
    cell_emp = table.cell(0, 0)
    if empresa and empresa.firma_representante:
        try:
            p = cell_emp.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(empresa.firma_representante.path, height=Inches(0.6))
        except Exception:
            pass

    cell_emp_name = table.cell(1, 0)
    p = cell_emp_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 30)
    run.font.size = Pt(9)

    cell_emp_info = table.cell(2, 0)
    p = cell_emp_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if empresa:
        run = p.add_run(f"{empresa.representante_legal or empresa.razon_social}")
        run.font.size = Pt(9)
        run.bold = True
        p.add_run('\n')
        run2 = p.add_run(f"{empresa.cargo_representante or 'Representante Legal'}\nEL EMPLEADOR")
        run2.font.size = Pt(8)

    # Worker signature
    cell_trab = table.cell(1, 2)
    p = cell_trab.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 30)
    run.font.size = Pt(9)

    cell_trab_info = table.cell(2, 2)
    p = cell_trab_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if personal:
        run = p.add_run(personal.apellidos_nombres or '---')
        run.font.size = Pt(9)
        run.bold = True
        p.add_run('\n')
        run2 = p.add_run(f"{personal.tipo_doc}: {personal.nro_doc}\nEL TRABAJADOR")
        run2.font.size = Pt(8)

    # ── Footer ──
    footer_section = section.footer
    footer_para = footer_section.paragraphs[0] if footer_section.paragraphs else footer_section.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if empresa:
        run = footer_para.add_run(
            f"{empresa.razon_social} | RUC {empresa.ruc} | Documento generado por Harmoni ERP"
        )
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(153, 153, 153)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
